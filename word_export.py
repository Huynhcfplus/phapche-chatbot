
from docx import Document
from docx.shared import Inches
from docxtpl import DocxTemplate

def export_to_word(context, flowchart_img_path, template_path='templates/quy_trinh_template.docx'):
    doc = DocxTemplate(template_path)
    doc.render(context)
    output_path = "quy_trinh_output.docx"
    doc.save(output_path)

    docx = Document(output_path)
    docx.add_page_break()
    docx.add_heading("Lưu đồ quy trình", level=2)
    docx.add_picture(flowchart_img_path, width=Inches(5.5))
    docx.save(output_path)
    return output_path
